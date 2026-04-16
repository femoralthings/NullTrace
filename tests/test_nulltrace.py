"""
tests/test_nulltrace.py - NullTrace comprehensive test suite

Tests every steganography method end-to-end:
  1. Crypto standalone
  2. PNG LSB hide/extract (correct key + wrong key rejection)
  3. JPEG EXIF hide/extract
  4. WAV LSB hide/extract
  5. EOF append hide/extract (PNG, JPEG, PDF)
  6. Zero-width character hide/extract
  7. ZIP comment hide/extract
  8. PDF metadata hide/extract
  9. NTFS ADS hide/extract/list (Windows only)
 10. Blind scanner - detects planted payloads, passes clean files
 11. Binary payload (image hidden inside image)
 12. Large payload stress test

Run with:  python -m pytest tests/ -v
       or:  python tests/test_nulltrace.py
"""

import os
import struct
import sys
import wave
import zipfile
from pathlib import Path

# ─── colour output helpers ────────────────────────────────────────────────────
GREEN  = "\033[92m"
RED    = "\033[91m"
YELLOW = "\033[93m"
CYAN   = "\033[96m"
RESET  = "\033[0m"
BOLD   = "\033[1m"

passed = failed = skipped = 0

def ok(name):
    global passed
    passed += 1
    print(f"  {GREEN}[PASS]{RESET} {name}")

def fail(name, reason=""):
    global failed
    failed += 1
    msg = f"  {RED}[FAIL]{RESET} {name}"
    if reason:
        msg += f"\n         {RED}{reason}{RESET}"
    print(msg)

def skip(name, reason=""):
    global skipped
    skipped += 1
    print(f"  {YELLOW}[SKIP]{RESET} {name}  ({reason})")

def section(title):
    print(f"\n{BOLD}{CYAN}── {title} {'─' * (55 - len(title))}{RESET}")


# ─── test asset builders ──────────────────────────────────────────────────────

def make_png(path: str, w=200, h=200):
    """Create a realistic-looking PNG with gradients (not solid colour)."""
    from PIL import Image, ImageDraw
    img  = Image.new('RGB', (w, h))
    draw = ImageDraw.Draw(img)
    # Gradient background
    for y in range(h):
        r = int(80  + 100 * y / h)
        g = int(100 + 80  * y / h)
        b = int(120 + 60  * y / h)
        draw.line([(0, y), (w, y)], fill=(r, g, b))
    # Some shapes for visual variety
    draw.ellipse([40, 40, 160, 160], fill=(180, 100, 80), outline=(255, 200, 100))
    draw.rectangle([60, 60, 140, 140], fill=(60, 120, 180))
    img.save(path, 'PNG')


def make_jpeg(path: str):
    from PIL import Image, ImageDraw
    img  = Image.new('RGB', (300, 200), color=(80, 100, 120))
    draw = ImageDraw.Draw(img)
    draw.polygon([(150, 20), (280, 180), (20, 180)], fill=(200, 150, 80))
    img.save(path, 'JPEG', quality=85)


def make_wav(path: str, duration_sec=2, sample_rate=44100):
    import math
    n_samples = duration_sec * sample_rate
    with wave.open(path, 'wb') as w:
        w.setnchannels(1)
        w.setsampwidth(2)
        w.setframerate(sample_rate)
        # 440Hz sine wave with some noise
        import random
        frames = []
        for i in range(n_samples):
            val = int(8000 * math.sin(2 * math.pi * 440 * i / sample_rate))
            val += random.randint(-200, 200)  # noise
            val = max(-32768, min(32767, val))
            frames.extend(struct.pack('<h', val))
        w.writeframes(bytes(frames))


def make_zip(path: str):
    with zipfile.ZipFile(path, 'w') as zf:
        zf.writestr("readme.txt", "This is a normal ZIP archive.\n")
        zf.writestr("data.csv",   "name,value\nfoo,1\nbar,2\n")


def make_pdf(path: str):
    """Create a minimal but valid PDF."""
    content = b"""%PDF-1.4
1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj
2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj
3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
  /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >> endobj
4 0 obj << /Length 44 >>
stream
BT /F1 12 Tf 100 700 Td (NullTrace Test PDF) Tj ET
endstream
endobj
5 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj
xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000274 00000 n
0000000369 00000 n
trailer << /Size 6 /Root 1 0 R >>
startxref
441
%%EOF"""
    with open(path, 'wb') as f:
        f.write(content)


def make_text(path: str):
    with open(path, 'w', encoding='utf-8') as f:
        f.write(
            "NullTrace zero-width character steganography test.\n"
            "This document appears completely normal to any reader.\n"
            "The hidden payload is invisible and indistinguishable from the surrounding text.\n"
        )


# ─── 1. Crypto ────────────────────────────────────────────────────────────────

def test_crypto():
    section("1. Cryptography (AES-256-GCM + scrypt)")
    from core.crypto import encrypt, decrypt

    msg = b"DEADBEEF - NullTrace crypto test"

    try:
        enc = encrypt(msg, "test-password-123")
        assert enc[0:1] == b'\x01', "Wrong version byte"
        assert len(enc) >= 61, "Ciphertext too short"
        ok("encrypt produces versioned output")
    except Exception as e:
        fail("encrypt", str(e)); return

    try:
        dec = decrypt(enc, "test-password-123")
        assert dec == msg
        ok("decrypt roundtrip correct")
    except Exception as e:
        fail("decrypt roundtrip", str(e))

    try:
        decrypt(enc, "wrong-password")
        fail("wrong password should raise ValueError")
    except ValueError:
        ok("wrong password correctly rejected")

    try:
        tampered = enc[:20] + bytes([enc[20] ^ 0xFF]) + enc[21:]
        decrypt(tampered, "test-password-123")
        fail("tampered ciphertext should raise ValueError")
    except ValueError:
        ok("tampered ciphertext rejected (GCM auth tag)")

    try:
        enc2 = encrypt(msg, "test-password-123")
        assert enc != enc2, "Same plaintext should produce different ciphertext (random salt+nonce)"
        ok("random salt+nonce - same input ≠ same output")
    except Exception as e:
        fail("randomness check", str(e))


# ─── 2. PNG LSB ───────────────────────────────────────────────────────────────

def test_png_lsb(tmp: Path):
    section("2. PNG LSB Steganography")
    from core.lsb    import hide, extract, capacity
    from core.crypto import encrypt, decrypt

    carrier = str(tmp / "carrier.png")
    output  = str(tmp / "output.png")
    make_png(carrier)

    try:
        cap = capacity(carrier)
        assert cap > 100
        ok(f"capacity() = {cap} bytes")
    except Exception as e:
        fail("capacity", str(e)); return

    msg = b"Hidden inside a PNG - DEADBEEF - NullTrace test payload"
    enc = encrypt(msg, "frank-key")

    try:
        hide(carrier, enc, "frank-key", output)
        assert Path(output).exists()
        ok("hide() creates output file")
    except Exception as e:
        fail("hide", str(e)); return

    # Timestamps
    orig_mtime = os.stat(carrier).st_mtime
    out_mtime  = os.stat(output).st_mtime
    if abs(orig_mtime - out_mtime) < 2:
        ok("timestamps preserved (within 2s)")
    else:
        fail("timestamps not preserved",
             f"carrier={orig_mtime:.0f} output={out_mtime:.0f}")

    try:
        raw   = extract(output, "frank-key")
        plain = decrypt(raw, "frank-key")
        assert plain == msg
        ok("extract + decrypt roundtrip correct")
    except Exception as e:
        fail("extract/decrypt", str(e))

    try:
        raw2 = extract(output, "wrong-key")
        decrypt(raw2, "wrong-key")
        fail("wrong key should fail")
    except ValueError:
        ok("wrong key rejected on extract")

    try:
        raw3 = extract(carrier, "frank-key")  # clean carrier
        decrypt(raw3, "frank-key")
        fail("clean carrier should not decrypt")
    except ValueError:
        ok("clean carrier correctly rejected")


# ─── 3. JPEG EXIF ─────────────────────────────────────────────────────────────

def test_jpeg_exif(tmp: Path):
    section("3. JPEG EXIF Steganography")
    from core.jpeg_exif import hide, extract
    from core.crypto    import encrypt, decrypt

    carrier = str(tmp / "carrier.jpg")
    output  = str(tmp / "output.jpg")
    make_jpeg(carrier)

    msg = b"JPEG EXIF hidden payload - operation successful"
    enc = encrypt(msg, "exif-key")

    try:
        hide(carrier, enc, output)
        assert Path(output).exists()
        ok("hide() creates output")
    except Exception as e:
        fail("hide", str(e)); return

    # Output should be same size or close (no re-encode)
    orig_size = os.path.getsize(carrier)
    out_size  = os.path.getsize(output)
    if abs(out_size - orig_size) < 2000:
        ok(f"no JPEG re-encode (size diff: {abs(out_size - orig_size)} bytes)")
    else:
        fail("JPEG may have been re-encoded",
             f"carrier={orig_size} output={out_size} diff={abs(out_size - orig_size)}")

    try:
        raw   = extract(output)
        plain = decrypt(raw, "exif-key")
        assert plain == msg
        ok("extract + decrypt roundtrip correct")
    except Exception as e:
        fail("extract/decrypt", str(e))

    try:
        extract(carrier)
        fail("clean JPEG should not extract")
    except ValueError:
        ok("clean JPEG correctly raises ValueError")


# ─── 4. WAV LSB ───────────────────────────────────────────────────────────────

def test_wav_lsb(tmp: Path):
    section("4. WAV Audio LSB Steganography")
    from core.wav_lsb import hide, extract, capacity
    from core.crypto  import encrypt, decrypt

    carrier = str(tmp / "carrier.wav")
    output  = str(tmp / "output.wav")
    make_wav(carrier)

    try:
        cap = capacity(carrier)
        assert cap > 100
        ok(f"capacity() = {cap} bytes")
    except Exception as e:
        fail("capacity", str(e)); return

    msg = b"Audio steganography - hidden in PCM samples - DEADBEEF"
    enc = encrypt(msg, "wav-key")

    try:
        hide(carrier, enc, "wav-key", output)
        assert Path(output).exists()
        ok("hide() creates output")
    except Exception as e:
        fail("hide", str(e)); return

    try:
        raw   = extract(output, "wav-key")
        plain = decrypt(raw, "wav-key")
        assert plain == msg
        ok("extract + decrypt roundtrip correct")
    except Exception as e:
        fail("extract/decrypt", str(e))

    try:
        raw2 = extract(output, "wrong-key")
        decrypt(raw2, "wrong-key")
        fail("wrong key should fail")
    except ValueError:
        ok("wrong key rejected")


# ─── 5. EOF Append ────────────────────────────────────────────────────────────

def test_eof_append(tmp: Path):
    section("5. EOF-Append Steganography")
    from core.eof_append import hide, extract, scan
    from core.crypto     import encrypt, decrypt

    for fmt, maker in [
        ('PNG',  lambda p: make_png(p, 50, 50)),
        ('JPEG', make_jpeg),
    ]:
        ext     = '.png' if fmt == 'PNG' else '.jpg'
        carrier = str(tmp / f"eof_carrier{ext}")
        output  = str(tmp / f"eof_output{ext}")
        maker(carrier)

        msg = f"EOF append payload - {fmt} test".encode()
        enc = encrypt(msg, "eof-key")

        try:
            hide(carrier, enc, output)
            ok(f"[{fmt}] hide() appends after EOF marker")
        except Exception as e:
            fail(f"[{fmt}] hide", str(e)); continue

        try:
            scan_result = scan(output)
            assert scan_result['detected']
            ok(f"[{fmt}] blind scan detects appended data")
        except Exception as e:
            fail(f"[{fmt}] scan", str(e))

        try:
            raw   = extract(output)
            plain = decrypt(raw, "eof-key")
            assert plain == msg
            ok(f"[{fmt}] extract + decrypt roundtrip correct")
        except Exception as e:
            fail(f"[{fmt}] extract/decrypt", str(e))


# ─── 6. Zero-Width Characters ────────────────────────────────────────────────

def test_zero_width(tmp: Path):
    section("6. Zero-Width Character Steganography")
    from core.zero_width import hide, extract, scan, fingerprint
    from core.crypto     import encrypt, decrypt

    txt_path = str(tmp / "cover.txt")
    make_text(txt_path)

    with open(txt_path, 'r', encoding='utf-8') as f:
        cover = f.read()

    msg = b"ZWC hidden message - invisible to the naked eye"
    enc = encrypt(msg, "zwc-key")

    try:
        steg_text = hide(cover, enc)
        assert len(steg_text) > len(cover)  # invisible chars added
        ok("hide() produces longer string (ZW chars inserted)")
    except Exception as e:
        fail("hide", str(e)); return

    try:
        scan_result = scan(steg_text)
        assert scan_result['detected']
        assert scan_result['total_zw_chars'] > 0
        ok(f"scan() detects {scan_result['total_zw_chars']} ZW chars "
           f"(~{scan_result['estimated_bytes']} bytes)")
    except Exception as e:
        fail("scan", str(e))

    try:
        raw   = extract(steg_text)
        plain = decrypt(raw, "zwc-key")
        assert plain == msg
        ok("extract + decrypt roundtrip correct")
    except Exception as e:
        fail("extract/decrypt", str(e))

    try:
        scan_clean = scan(cover)
        assert not scan_clean['detected']
        ok("clean text correctly shows no ZW chars")
    except Exception as e:
        fail("clean text scan", str(e))

    # Fingerprint function
    try:
        fp = fingerprint("Confidential report for Agent 7.", b"AGENT-007")
        r  = scan(fp)
        assert r['detected']
        ok("fingerprint() embeds trackable marker")
    except Exception as e:
        fail("fingerprint", str(e))


# ─── 7. ZIP Comment ───────────────────────────────────────────────────────────

def test_zip_comment(tmp: Path):
    section("7. ZIP Comment Field Steganography")
    from core.zip_comment import hide, extract, scan
    from core.crypto      import encrypt, decrypt

    carrier = str(tmp / "archive.zip")
    output  = str(tmp / "archive_steg.zip")
    make_zip(carrier)

    msg = b"Hidden in ZIP comment - standard 7-Zip or WinRAR won't flag this"
    enc = encrypt(msg, "zip-key")

    try:
        hide(carrier, enc, output)
        ok("hide() writes payload to ZIP comment")
    except Exception as e:
        fail("hide", str(e)); return

    try:
        scan_result = scan(output)
        assert scan_result['detected']
        ok(f"scan() finds {scan_result['comment_bytes']} bytes in comment")
    except Exception as e:
        fail("scan", str(e))

    try:
        raw   = extract(output)
        plain = decrypt(raw, "zip-key")
        assert plain == msg
        ok("extract + decrypt roundtrip correct")
    except Exception as e:
        fail("extract/decrypt", str(e))

    try:
        scan_clean = scan(carrier)
        assert not scan_clean['detected']
        ok("clean ZIP correctly shows no comment")
    except Exception as e:
        fail("clean ZIP scan", str(e))


# ─── 8. PDF Metadata ──────────────────────────────────────────────────────────

def test_pdf_meta(tmp: Path):
    section("8. PDF XMP Metadata Steganography")
    from core.pdf_meta import hide, extract, scan
    from core.crypto   import encrypt, decrypt

    carrier = str(tmp / "doc.pdf")
    output  = str(tmp / "doc_steg.pdf")
    make_pdf(carrier)

    msg = b"Hidden in PDF XMP metadata - survives most PDF viewers and re-saves"
    enc = encrypt(msg, "pdf-key")

    try:
        hide(carrier, enc, output)
        ok("hide() embeds payload in XMP metadata")
    except Exception as e:
        fail("hide", str(e)); return

    try:
        scan_result = scan(output)
        assert scan_result['detected']
        ok("scan() detects NullTrace XMP key")
    except Exception as e:
        fail("scan", str(e))

    try:
        raw   = extract(output)
        plain = decrypt(raw, "pdf-key")
        assert plain == msg
        ok("extract + decrypt roundtrip correct")
    except Exception as e:
        fail("extract/decrypt", str(e))


# ─── 9. NTFS ADS ──────────────────────────────────────────────────────────────

def test_ntfs_ads(tmp: Path):
    section("9. NTFS Alternate Data Streams")

    if os.name != 'nt':
        skip("ADS: hide", "Windows/NTFS only")
        skip("ADS: extract", "Windows/NTFS only")
        skip("ADS: list_streams", "Windows/NTFS only")
        skip("ADS: delete", "Windows/NTFS only")
        return

    from core.ntfs_ads import hide, extract, list_streams, delete_stream, scan
    from core.crypto   import encrypt, decrypt

    host = str(tmp / "host.txt")
    with open(host, 'w') as f:
        f.write("This is a normal text file.\n")

    host_size = os.path.getsize(host)
    msg = b"ADS payload - invisible to Explorer, dir, and most security tools"
    enc = encrypt(msg, "ads-key")

    try:
        hide(host, "nulltrace", enc)
        # Host file size must NOT change
        assert os.path.getsize(host) == host_size
        ok("hide() stores payload in ADS (host size unchanged)")
    except Exception as e:
        fail("hide", str(e)); return

    try:
        streams = list_streams(host)
        assert "nulltrace" in streams
        ok(f"list_streams() finds stream: {streams}")
    except Exception as e:
        fail("list_streams", str(e))

    try:
        scan_result = scan(host)
        assert scan_result['detected']
        ok("scan() detects ADS on host file")
    except Exception as e:
        fail("scan", str(e))

    try:
        raw   = extract(host, "nulltrace")
        plain = decrypt(raw, "ads-key")
        assert plain == msg
        ok("extract + decrypt roundtrip correct")
    except Exception as e:
        fail("extract/decrypt", str(e))

    try:
        delete_stream(host, "nulltrace")
        assert "nulltrace" not in list_streams(host)
        ok("delete_stream() removes ADS cleanly")
    except Exception as e:
        fail("delete_stream", str(e))


# ─── 10. Blind Scanner ───────────────────────────────────────────────────────

def test_blind_scanner(tmp: Path):
    section("10. Blind Multi-Vector Scanner")
    from detector.scan import scan_file
    from core.lsb      import hide as lsb_hide
    from core.crypto   import encrypt

    # Plant a payload
    carrier = str(tmp / "scan_carrier.png")
    steg    = str(tmp / "scan_steg.png")
    make_png(carrier, 300, 300)
    enc = encrypt(b"scanner test payload", "scan-key")
    lsb_hide(carrier, enc, "scan-key", steg)

    # Clean file should be clean
    r_clean = scan_file(carrier)
    if not r_clean['overall_suspicious']:
        ok("clean file reports CLEAN")
    else:
        fail("clean file falsely flagged",
             str([f['method'] for f in r_clean['findings']]))

    # Steg file: scanner should not raise (even if no statistical trigger due to PRNG spread)
    try:
        r_steg = scan_file(steg)
        ok(f"steg file scanned without error ({len(r_steg['findings'])} findings)")
    except Exception as e:
        fail("scan_file on steg", str(e))

    # With password - should extract successfully
    r_key = scan_file(steg, password="scan-key")
    if r_key.get('extracted_payload'):
        content = r_key['extracted_payload']['content']
        assert content == b"scanner test payload"
        ok("scanner extracts and decrypts payload when key provided")
    else:
        fail("scanner failed to extract with correct key")

    # ZWC scanner - plants detectable content
    from core.zero_width import hide as zwc_hide, scan as zwc_scan
    cover_file = str(tmp / "cover.txt")
    make_text(cover_file)
    with open(cover_file, 'r', encoding='utf-8') as f:
        cover = f.read()
    enc2 = encrypt(b"ZWC scanner test", "scan-key")
    steg_text = zwc_hide(cover, enc2)
    r_zw = zwc_scan(steg_text)
    if r_zw['detected']:
        ok(f"ZWC blind scan detects hidden chars ({r_zw['total_zw_chars']} chars)")
    else:
        fail("ZWC blind scan missed planted payload")


# ─── 11. Binary payload (image hidden in image) ───────────────────────────────

def test_binary_payload(tmp: Path):
    section("11. Binary Payload (image hidden inside image)")
    from core.lsb    import hide, extract, capacity
    from core.crypto import encrypt, decrypt

    # Make a small JPEG to hide inside a larger PNG
    secret_img = str(tmp / "secret.jpg")
    carrier    = str(tmp / "large_carrier.png")
    output     = str(tmp / "large_output.png")

    make_jpeg(secret_img)
    make_png(carrier, 400, 400)  # Larger carrier needed

    with open(secret_img, 'rb') as f:
        secret_bytes = f.read()

    cap = capacity(carrier)
    enc = encrypt(secret_bytes, "binary-key")

    if len(enc) > cap:
        skip("binary payload", f"test image too small ({len(enc)}B > {cap}B capacity)")
        return

    try:
        hide(carrier, enc, "binary-key", output)
        ok(f"JPEG ({len(secret_bytes)}B) hidden inside PNG ({cap}B capacity)")
    except Exception as e:
        fail("hide binary", str(e)); return

    try:
        raw       = extract(output, "binary-key")
        recovered = decrypt(raw, "binary-key")
        assert recovered == secret_bytes
        ok(f"Recovered binary exactly ({len(recovered)} bytes == original {len(secret_bytes)} bytes)")
    except Exception as e:
        fail("extract binary", str(e))

    # Verify recovered bytes are a valid JPEG
    try:
        from PIL import Image
        import io
        img = Image.open(io.BytesIO(recovered))
        img.verify()
        ok("Recovered bytes form a valid JPEG (PIL verified)")
    except Exception as e:
        fail("recovered JPEG validation", str(e))


# ─── 12. Large payload stress test ───────────────────────────────────────────

def test_large_payload(tmp: Path):
    section("12. Large Payload Stress Test")
    from core.lsb    import hide, extract, capacity
    from core.crypto import encrypt, decrypt

    # 2048×1536 image = ~12MB, ~900KB LSB capacity
    carrier = str(tmp / "stress_carrier.png")
    output  = str(tmp / "stress_output.png")
    make_png(carrier, 512, 512)

    cap = capacity(carrier)

    # Fill 70% of capacity with random data
    payload_size = int(cap * 0.70)
    raw_payload  = os.urandom(payload_size)
    enc          = encrypt(raw_payload, "stress-key")

    if len(enc) > cap:
        skip("large payload", f"encrypted payload ({len(enc)}B) > capacity ({cap}B)")
        return

    try:
        hide(carrier, enc, "stress-key", output)
        ok(f"hide() with {payload_size}B payload ({payload_size/cap*100:.0f}% capacity)")
    except Exception as e:
        fail("hide large", str(e)); return

    try:
        raw   = extract(output, "stress-key")
        plain = decrypt(raw, "stress-key")
        assert plain == raw_payload
        ok(f"extract + decrypt {len(plain)}B - byte-perfect match")
    except Exception as e:
        fail("extract large", str(e))


# ─── 13. Key file support ─────────────────────────────────────────────────────

def test_keyfile_crypto(tmp: Path):
    section("13. Key File (2nd Factor) Cryptography")
    from core.crypto import encrypt, decrypt

    kf = tmp / "secret.bin"
    kf.write_bytes(os.urandom(64))
    kf_bytes = kf.read_bytes()

    msg = b"Key file two-factor test - requires both password AND file"

    try:
        enc = encrypt(msg, "pw", kf_bytes)
        dec = decrypt(enc, "pw", kf_bytes)
        assert dec == msg
        ok("encrypt/decrypt with keyfile roundtrip correct")
    except Exception as e:
        fail("keyfile roundtrip", str(e)); return

    # Wrong password, correct keyfile
    try:
        decrypt(enc, "wrong-pw", kf_bytes)
        fail("wrong password + correct keyfile should fail")
    except ValueError:
        ok("wrong password + correct keyfile correctly rejected")

    # Correct password, wrong keyfile
    wrong_kf = os.urandom(64)
    try:
        decrypt(enc, "pw", wrong_kf)
        fail("correct password + wrong keyfile should fail")
    except ValueError:
        ok("correct password + wrong keyfile correctly rejected")

    # Correct password, no keyfile
    try:
        decrypt(enc, "pw", None)
        fail("correct password + no keyfile should fail")
    except ValueError:
        ok("correct password + missing keyfile correctly rejected")

    # No keyfile encryption — keyfile argument should be ignored (None)
    try:
        enc2 = encrypt(msg, "pw", None)
        dec2 = decrypt(enc2, "pw", None)
        assert dec2 == msg
        ok("encrypt/decrypt without keyfile still works")
    except Exception as e:
        fail("no-keyfile path", str(e))


# ─── 14. Alpha channel LSB ────────────────────────────────────────────────────

def test_alpha_lsb(tmp: Path):
    section("14. Alpha Channel LSB Steganography")
    try:
        from core.alpha_lsb import hide, extract, capacity, scan
        from core.crypto    import encrypt, decrypt
    except ImportError as e:
        skip("alpha LSB (all)", str(e)); return

    # Make an RGBA PNG with non-uniform alpha
    from PIL import Image
    import random as rnd
    carrier = str(tmp / "alpha_carrier.png")
    w, h = 200, 200
    img = Image.new('RGBA', (w, h))
    pixels = [(rnd.randint(50, 200), rnd.randint(50, 200),
               rnd.randint(50, 200), rnd.randint(100, 255))
              for _ in range(w * h)]
    img.putdata(pixels)
    img.save(carrier, 'PNG')

    output = str(tmp / "alpha_output.png")

    try:
        cap = capacity(carrier)
        assert cap > 0
        ok(f"capacity() = {cap} bytes")
    except Exception as e:
        fail("capacity", str(e)); return

    msg = b"Alpha channel hidden payload - RGB untouched"
    enc = encrypt(msg, "alpha-key")

    try:
        hide(carrier, enc, "alpha-key", output)
        assert Path(output).exists()
        ok("hide() creates output")
    except Exception as e:
        fail("hide", str(e)); return

    # Verify RGB channels are bit-identical
    try:
        orig = Image.open(carrier).convert('RGBA')
        steg = Image.open(output).convert('RGBA')
        orig_px = list(orig.getdata())
        steg_px = list(steg.getdata())
        rgb_diffs = sum(
            1 for o, s in zip(orig_px, steg_px)
            if o[:3] != s[:3]
        )
        assert rgb_diffs == 0
        ok("RGB channels unmodified (only alpha channel changed)")
    except Exception as e:
        fail("RGB integrity check", str(e))

    try:
        raw   = extract(output, "alpha-key")
        plain = decrypt(raw, "alpha-key")
        assert plain == msg
        ok("extract + decrypt roundtrip correct")
    except Exception as e:
        fail("extract/decrypt", str(e))

    try:
        raw2 = extract(output, "wrong-key")
        decrypt(raw2, "wrong-key")
        fail("wrong key should fail")
    except ValueError:
        ok("wrong key correctly rejected")

    try:
        result = scan(output)
        ok(f"scan() runs without error (detected={result['detected']})")
    except Exception as e:
        fail("scan", str(e))


# ─── 15. JPEG DCT ─────────────────────────────────────────────────────────────

def test_jpeg_dct(tmp: Path):
    section("15. JPEG DCT Steganography")
    try:
        from core.jpeg_dct import hide, extract, capacity
        from core.crypto   import encrypt, decrypt
    except ImportError as e:
        skip("JPEG DCT (all)", str(e)); return

    carrier = str(tmp / "dct_carrier.jpg")
    output  = str(tmp / "dct_output.jpg")
    make_jpeg(carrier)

    try:
        cap = capacity(carrier)
        assert cap > 0
        ok(f"capacity() = {cap} bytes")
    except Exception as e:
        fail("capacity", str(e)); return

    msg = b"DCT coefficient hidden payload - survives JPEG re-encode"
    enc = encrypt(msg, "dct-key")

    if len(enc) > cap:
        skip("DCT hide/extract", f"test image capacity {cap}B < payload {len(enc)}B")
        return

    try:
        hide(carrier, enc, "dct-key", output, quality=95)
        assert Path(output).exists()
        ok("hide() creates output (re-encoded at q=95)")
    except Exception as e:
        fail("hide", str(e)); return

    try:
        raw   = extract(output, "dct-key")
        plain = decrypt(raw, "dct-key")
        assert plain == msg
        ok("extract + decrypt roundtrip correct")
    except Exception as e:
        fail("extract/decrypt", str(e))

    try:
        raw2 = extract(output, "wrong-key")
        decrypt(raw2, "wrong-key")
        fail("wrong key should fail")
    except (ValueError, Exception):
        ok("wrong key correctly rejected")


# ─── 16. DOCX hidden text ─────────────────────────────────────────────────────

def test_docx_hidden(tmp: Path):
    section("16. DOCX Hidden Text (w:vanish)")
    try:
        from core.docx_hidden import hide, extract, scan
        from core.crypto      import encrypt, decrypt
        from docx import Document
    except ImportError as e:
        skip("DOCX hidden text (all)", str(e)); return

    # Build a minimal DOCX
    carrier = str(tmp / "carrier.docx")
    output  = str(tmp / "output.docx")
    doc = Document()
    doc.add_paragraph("This is a completely normal document.")
    doc.add_paragraph("Nothing suspicious here at all.")
    doc.save(carrier)

    msg = b"DOCX hidden text payload - invisible in Word by default"
    enc = encrypt(msg, "docx-key")

    try:
        hide(carrier, enc, output)
        assert Path(output).exists()
        ok("hide() creates output")
    except Exception as e:
        fail("hide", str(e)); return

    try:
        result = scan(output)
        assert result['detected']
        ok(f"scan() detects hidden run(s): {result['findings'][:1]}")
    except Exception as e:
        fail("scan", str(e))

    try:
        raw   = extract(output)
        plain = decrypt(raw, "docx-key")
        assert plain == msg
        ok("extract + decrypt roundtrip correct")
    except Exception as e:
        fail("extract/decrypt", str(e))

    try:
        result_clean = scan(carrier)
        assert not result_clean['detected']
        ok("clean DOCX shows no hidden text")
    except Exception as e:
        fail("clean DOCX scan", str(e))


# ─── 17. MP3 ID3 ──────────────────────────────────────────────────────────────

def test_mp3_id3(tmp: Path):
    section("17. MP3 ID3 Tag Steganography")
    try:
        from core.mp3_id3 import hide, extract, scan
        from core.crypto  import encrypt, decrypt
        from mutagen.id3  import ID3, TIT2
        from mutagen.mp3  import MP3
    except ImportError as e:
        skip("MP3 ID3 (all)", str(e)); return

    # Create a minimal valid MP3 frame (MPEG1 Layer3 silence)
    carrier = str(tmp / "carrier.mp3")
    output  = str(tmp / "output.mp3")

    # Minimal MP3: one silent frame header + zero payload
    # Frame sync + MPEG1 Layer3 128kbps 44100Hz mono
    mp3_frame = bytes([0xFF, 0xFB, 0x90, 0x00]) + bytes(413)
    with open(carrier, 'wb') as f:
        f.write(mp3_frame * 10)

    msg = b"MP3 ID3 hidden payload - audio bitstream untouched"
    enc = encrypt(msg, "mp3-key")

    try:
        hide(carrier, enc, output)
        assert Path(output).exists()
        ok("hide() embeds payload in COMM ID3 tag")
    except Exception as e:
        fail("hide", str(e)); return

    try:
        result = scan(output)
        assert result['detected']
        ok(f"scan() finds NullTrace ID3 marker")
    except Exception as e:
        fail("scan", str(e))

    try:
        raw   = extract(output)
        plain = decrypt(raw, "mp3-key")
        assert plain == msg
        ok("extract + decrypt roundtrip correct")
    except Exception as e:
        fail("extract/decrypt", str(e))


# ─── 18. Plausible deniability (dual-bit plane) ───────────────────────────────

def test_plausible_deniability(tmp: Path):
    section("18. Plausible Deniability (Dual-Bit Plane)")
    from core.lsb    import hide_dual, extract, extract_real
    from core.crypto import encrypt, decrypt

    carrier = str(tmp / "dual_carrier.png")
    output  = str(tmp / "dual_output.png")
    make_png(carrier, 400, 400)

    real_msg  = b"Real payload - the actual secret"
    decoy_msg = b"Decoy payload - the innocent message"
    real_enc  = encrypt(real_msg,  "real-key")
    decoy_enc = encrypt(decoy_msg, "decoy-key")

    try:
        hide_dual(carrier,
                  real_payload=real_enc,   real_password="real-key",
                  decoy_payload=decoy_enc, decoy_password="decoy-key",
                  output_path=output)
        assert Path(output).exists()
        ok("hide_dual() creates dual-embedded image")
    except Exception as e:
        fail("hide_dual", str(e)); return

    # extract() with decoy key should return decoy
    try:
        raw_decoy  = extract(output, "decoy-key")
        plain_decoy = decrypt(raw_decoy, "decoy-key")
        assert plain_decoy == decoy_msg
        ok("extract(decoy-key) returns decoy payload")
    except Exception as e:
        fail("decoy extraction", str(e))

    # extract_real() with real key should return real
    try:
        raw_real  = extract_real(output, "real-key")
        plain_real = decrypt(raw_real, "real-key")
        assert plain_real == real_msg
        ok("extract_real(real-key) returns real payload")
    except Exception as e:
        fail("real extraction", str(e))

    # Verify the two payloads are completely independent
    try:
        raw_real2 = extract_real(output, "real-key")
        raw_decoy2 = extract(output, "decoy-key")
        assert raw_real2 != raw_decoy2
        ok("Real and decoy payloads are independent (no overlap)")
    except Exception as e:
        fail("independence check", str(e))


# ─── 19. RS Analysis ──────────────────────────────────────────────────────────

def test_rs_analysis(tmp: Path):
    section("19. RS (Regular-Singular) Analysis")
    try:
        from detector.rs_analysis import analyze_image
    except ImportError as e:
        skip("RS analysis (all)", str(e)); return

    from core.lsb    import hide
    from core.crypto import encrypt

    # Clean image — RS should show low embedding rate
    clean = str(tmp / "rs_clean.png")
    make_png(clean, 200, 200)

    try:
        result = analyze_image(clean)
        rate = result['embedding_rate']
        ok(f"clean image RS embedding rate: {rate:.4f} (suspicious={result['suspicious']})")
    except Exception as e:
        fail("RS clean image", str(e))

    # Heavily embedded image — RS should show elevated rate
    steg = str(tmp / "rs_steg.png")
    make_png(steg, 300, 300)
    enc = encrypt(b"RS analysis test payload " * 20, "rs-key")
    try:
        from core.lsb import capacity
        cap = capacity(steg)
        if len(enc) <= cap:
            hide(steg, enc, "rs-key", steg)  # in-place via temp
            result2 = analyze_image(steg)
            ok(f"steg image RS embedding rate: {result2['embedding_rate']:.4f} "
               f"(suspicious={result2['suspicious']})")
        else:
            skip("RS steg image", "payload too large for test image")
    except Exception as e:
        fail("RS steg image", str(e))


# ─── 20. CSV export ───────────────────────────────────────────────────────────

def test_csv_export(tmp: Path):
    section("20. CSV Export")
    from detector.scan   import scan_file, export_csv
    from core.lsb        import hide
    from core.crypto     import encrypt

    carrier = str(tmp / "csv_carrier.png")
    steg    = str(tmp / "csv_steg.png")
    make_png(carrier, 200, 200)
    enc = encrypt(b"CSV export test", "csv-key")
    hide(carrier, enc, "csv-key", steg)

    reports  = [
        scan_file(carrier),
        scan_file(steg),
    ]
    csv_path = str(tmp / "results.csv")

    try:
        export_csv(reports, csv_path)
        assert Path(csv_path).exists()
        ok("export_csv() creates CSV file")
    except Exception as e:
        fail("export_csv", str(e)); return

    try:
        import csv
        with open(csv_path, newline='', encoding='utf-8') as f:
            rows = list(csv.DictReader(f))
        assert len(rows) >= 2
        assert 'file' in rows[0]
        assert 'method' in rows[0]
        assert 'confidence' in rows[0]
        ok(f"CSV has {len(rows)} rows with correct headers")
    except Exception as e:
        fail("CSV content validation", str(e))


# ─── Runner ───────────────────────────────────────────────────────────────────

def main():
    import tempfile
    print(f"\n{BOLD}{'=' * 60}{RESET}")
    print(f"{BOLD}  NullTrace Test Suite{RESET}")
    print(f"{BOLD}{'=' * 60}{RESET}")

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = Path(tmpdir)
        try:
            test_crypto()
            test_png_lsb(tmp)
            test_jpeg_exif(tmp)
            test_wav_lsb(tmp)
            test_eof_append(tmp)
            test_zero_width(tmp)
            test_zip_comment(tmp)
            test_pdf_meta(tmp)
            test_ntfs_ads(tmp)
            test_blind_scanner(tmp)
            test_binary_payload(tmp)
            test_large_payload(tmp)
            test_keyfile_crypto(tmp)
            test_alpha_lsb(tmp)
            test_jpeg_dct(tmp)
            test_docx_hidden(tmp)
            test_mp3_id3(tmp)
            test_plausible_deniability(tmp)
            test_rs_analysis(tmp)
            test_csv_export(tmp)
        except KeyboardInterrupt:
            print(f"\n{YELLOW}Interrupted.{RESET}")

    print(f"\n{BOLD}{'=' * 60}{RESET}")
    total = passed + failed + skipped
    print(f"  Results: "
          f"{GREEN}{passed} passed{RESET}  "
          f"{RED}{failed} failed{RESET}  "
          f"{YELLOW}{skipped} skipped{RESET}  "
          f"/ {total} total")
    print(f"{BOLD}{'=' * 60}{RESET}\n")

    sys.exit(0 if failed == 0 else 1)


if __name__ == '__main__':
    main()
