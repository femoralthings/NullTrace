"""
core/docx_hidden.py — DOCX hidden text steganography

Microsoft Word supports a "Hidden" character format (w:vanish XML element).
Text marked as hidden:
  - Is invisible in normal document view
  - Does NOT appear when printing
  - IS visible only when "Show hidden text" is enabled in Word settings
  - Survives email, cloud sync, and most PDF export (unless "print hidden")
  - Is stored in plaintext XML inside the .docx ZIP — readable by NullTrace

We encode the payload as a hex string and insert it as a hidden run inside
the first paragraph of the document. The hex string is flanked by a NullTrace
sentinel so we can locate and extract it precisely.

Detection: scan all paragraph runs for w:vanish formatting. Hidden runs
with base64/hex content are a strong indicator of steganographic content.
"""

import os
import shutil
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml   import OxmlElement

NT_SENTINEL_OPEN  = "NT{"
NT_SENTINEL_CLOSE = "}TN"


def _make_hidden_run(paragraph, text: str) -> None:
    """Add a run with w:vanish (hidden) formatting to a paragraph."""
    run = paragraph.add_run(text)
    # Add w:vanish to the run's rPr (run properties)
    rPr = run._r.get_or_add_rPr()
    vanish = OxmlElement('w:vanish')
    rPr.append(vanish)
    # Set font colour to white as a belt-and-suspenders measure
    color = OxmlElement('w:color')
    color.set(qn('w:val'), 'FFFFFF')
    rPr.append(color)


def hide(docx_path: str, payload: bytes, output_path: str) -> None:
    """
    Embed payload as a hidden run in the first paragraph of the document.
    The payload is hex-encoded and wrapped in a NullTrace sentinel.
    """
    stat = os.stat(docx_path)
    original_times = (stat.st_atime, stat.st_mtime)

    shutil.copy2(docx_path, output_path)
    doc = Document(output_path)

    # Find or create first paragraph
    if not doc.paragraphs:
        doc.add_paragraph()

    para  = doc.paragraphs[0]
    token = NT_SENTINEL_OPEN + payload.hex() + NT_SENTINEL_CLOSE
    _make_hidden_run(para, token)

    doc.save(output_path)
    os.utime(output_path, original_times)


def extract(docx_path: str) -> bytes:
    """
    Find and extract the NullTrace hidden payload from a DOCX file.
    Returns raw (still-encrypted) bytes.
    """
    doc = Document(docx_path)

    full_text = ''
    for para in doc.paragraphs:
        for run in para.runs:
            rPr = run._r.find(qn('w:rPr'))
            if rPr is not None:
                vanish = rPr.find(qn('w:vanish'))
                if vanish is not None:
                    full_text += run.text

    start = full_text.find(NT_SENTINEL_OPEN)
    end   = full_text.find(NT_SENTINEL_CLOSE)

    if start == -1 or end == -1:
        raise ValueError("No NullTrace payload found in DOCX hidden text.")

    hex_payload = full_text[start + len(NT_SENTINEL_OPEN):end]

    try:
        return bytes.fromhex(hex_payload)
    except ValueError:
        raise ValueError("DOCX hidden payload is malformed (not valid hex).")


def scan(docx_path: str) -> dict:
    """
    Scan a DOCX for hidden text runs (w:vanish), white-on-white text,
    and other steganographic indicators.
    """
    result = {'detected': False, 'findings': []}

    try:
        doc = Document(docx_path)
        hidden_runs   = 0
        white_runs    = 0
        nulltrace_hit = False
        tiny_runs     = 0

        for para in doc.paragraphs:
            for run in para.runs:
                rPr = run._r.find(qn('w:rPr'))
                if rPr is None:
                    continue

                # Check w:vanish
                if rPr.find(qn('w:vanish')) is not None:
                    hidden_runs += 1
                    if NT_SENTINEL_OPEN in run.text:
                        nulltrace_hit = True

                # Check white colour
                color_el = rPr.find(qn('w:color'))
                if color_el is not None:
                    val = color_el.get(qn('w:val'), '')
                    if val.upper() in ('FFFFFF', 'FFFFFE', 'WHITE'):
                        white_runs += 1

                # Check tiny font size (< 2pt)
                sz = rPr.find(qn('w:sz'))
                if sz is not None:
                    try:
                        if int(sz.get(qn('w:val'), '24')) < 4:
                            tiny_runs += 1
                    except ValueError:
                        pass

        if nulltrace_hit:
            result['detected'] = True
            result['findings'].append("NullTrace hidden payload sentinel found")

        if hidden_runs:
            result['detected'] = True
            result['findings'].append(
                f"{hidden_runs} hidden run(s) with w:vanish formatting"
            )

        if white_runs:
            result['detected'] = True
            result['findings'].append(
                f"{white_runs} run(s) with white (#FFFFFF) text colour"
            )

        if tiny_runs:
            result['detected'] = True
            result['findings'].append(
                f"{tiny_runs} run(s) with font size < 2pt"
            )

        # Check document metadata for anomalies
        core = doc.core_properties
        if core.revision and int(str(core.revision)) > 50:
            result['findings'].append(
                f"High revision count: {core.revision} — many edit cycles"
            )

    except Exception as e:
        result['findings'].append(f"Error scanning DOCX: {e}")

    return result
